from docx import Document

document = Document('new_test.docx')
#插入6级标题
title = document.add_heading('this is title',0)
title = document.add_heading('this is title1',1)
title = document.add_heading('this is title2',2)
title = document.add_heading('this is title3',3)
title = document.add_heading('this is title4',4)
title = document.add_heading('this is title5',5)
title = document.add_heading('this is title6',6)
#插入文本
paragraph = document.add_paragraph('this is paragraph')
#分页
document.add_page_break()
#表格
records = (
    (3, '101', 'Spam'),
    (7, '422', 'Eggs'),
    (4, '631', 'Spam, spam, eggs, and spam')
)

tables = document.add_table(rows=1, cols=3)


#另存为
document.save('new_test1.docx')